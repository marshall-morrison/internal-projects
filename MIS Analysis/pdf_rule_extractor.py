#!/usr/bin/env python3
"""
PDF Rule Extractor - Analyze 150+ PDF documents to extract generalizable rules
for merchant instructions and context documents.

This script processes PDFs in batches, extracts text, analyzes patterns,
and synthesizes generalizable rules with merchant-specific customizations.
"""

import os
import json
import logging
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import hashlib
import time

from docx import Document
# PDF processing
import PyPDF2
import pdfplumber
from pdfminer.high_level import extract_text

# Text processing
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from collections import Counter

# LLM integration
import openai
from anthropic import Anthropic

# Data handling
import pandas as pd
import numpy as np

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    try:
        nltk.download('punkt_tab')
    except:
        print("Warning: Could not download punkt_tab. Using basic tokenization.")

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    try:
        nltk.download('stopwords')
    except:
        print("Warning: Could not download stopwords. Using empty stopwords list.")

@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""
    document_id: str
    chunk_id: str
    content: str
    chunk_type: str  # 'instructions', 'context', 'rules', 'general'
    page_number: Optional[int] = None
    word_count: int = 0

@dataclass
class AnalysisResult:
    """Results from LLM analysis of document chunks."""
    chunk_id: str
    document_id: str
    themes: List[str]
    rules: List[str]
    exceptions: List[str]
    merchant_specific: List[str]
    confidence_score: float
    analysis_timestamp: str

@dataclass
class SynthesizedRule:
    """A synthesized rule from multiple documents."""
    rule_id: str
    general_rule: str
    merchant_customizations: Dict[str, str]
    frequency: int
    confidence: float
    source_documents: List[str]

class PDFRuleExtractor:
    """Main class for extracting and analyzing rules from PDF documents."""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.setup_logging()
        self.setup_llm_clients()
        self.documents = []
        self.chunks = []
        self.analysis_results = []
        self.synthesized_rules = []
        self.raw_responses = {}  # Store raw LLM responses by document
        
    def setup_logging(self):
        """Setup logging configuration."""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('pdf_analysis.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        
    def setup_llm_clients(self):
        """Initialize LLM clients based on configuration."""
        self.openai_client = None
        self.anthropic_client = None
        
        # Prioritize Anthropic Claude
        if self.config.get('anthropic_api_key'):
            self.anthropic_client = Anthropic(api_key=self.config['anthropic_api_key'])
            self.logger.info("Anthropic Claude client configured successfully")
        else:
            self.logger.warning("No Anthropic API key found in configuration")
            
        # Fallback to OpenAI
        if self.config.get('openai_api_key'):
            openai.api_key = self.config['openai_api_key']
            self.openai_client = openai
            self.logger.info("OpenAI client configured as fallback")
            
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from a PDF file using multiple methods for robustness."""
        text = ""
        
        # Method 1: pdfplumber (best for complex layouts)
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    # Also try extracting text from tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " ".join([cell for cell in row if cell]) + "\n"
        except Exception as e:
            self.logger.warning(f"pdfplumber failed for {pdf_path}: {e}")
            
        # Method 2: PyPDF2 (fallback)
        if not text.strip():
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                self.logger.warning(f"PyPDF2 failed for {pdf_path}: {e}")
                
        # Method 3: pdfminer (final fallback)
        if not text.strip():
            try:
                text = extract_text(pdf_path)
            except Exception as e:
                self.logger.error(f"All PDF extraction methods failed for {pdf_path}: {e}")
                
        # Debug: Print first 200 characters if text was extracted
        if text.strip():
            self.logger.info(f"Extracted {len(text)} characters from {pdf_path}")
            self.logger.debug(f"First 200 chars: {text[:200]}")
        else:
            self.logger.warning(f"No text extracted from {pdf_path} - may be image-based PDF")
                
        return text.strip()
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(docx_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts)

            if text.strip():
                self.logger.info(f"Extracted {len(text)} characters from {docx_path}")
                self.logger.debug(f"First 200 chars: {text[:200]}")
            else:
                self.logger.warning(f"No text extracted from {docx_path}")

            return text.strip()

        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {docx_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from PDF or DOCX file based on file extension."""
        file_path_obj = Path(file_path)

        if file_path_obj.suffix.lower() == ".docx":
            return self.extract_text_from_docx(file_path)
        elif file_path_obj.suffix.lower() == ".pdf":
            return self.extract_text_from_pdf(file_path)
        else:
            self.logger.warning(f"Unsupported file type: {file_path_obj.suffix}")
            return ""
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Clean up common PDF artifacts
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\']+', ' ', text)
        
        return text.strip()
    
    def chunk_document(self, document_id: str, text: str) -> List[DocumentChunk]:
        """Split document into logical chunks for analysis."""
        chunks = []
        
        # Split by common section headers
        sections = re.split(r'\n(?=\s*(?:Instructions?|Rules?|Guidelines?|Context|Overview|Requirements?|Procedures?))', text, flags=re.IGNORECASE)
        
        chunk_id = 0
        for section in sections:
            if len(section.strip()) < 50:  # Skip very short sections
                continue
                
            # Determine chunk type based on content
            chunk_type = self.classify_chunk_type(section)
            
            # Further split large sections
            try:
                sentences = sent_tokenize(section)
            except:
                # Fallback to simple sentence splitting
                sentences = re.split(r'[.!?]+', section)
                sentences = [s.strip() for s in sentences if s.strip()]
            max_sentences = self.config.get('max_sentences_per_chunk', 10)
            
            for i in range(0, len(sentences), max_sentences):
                chunk_sentences = sentences[i:i + max_sentences]
                chunk_content = ' '.join(chunk_sentences)
                
                if len(chunk_content.strip()) < 100:  # Skip very short chunks
                    continue
                    
                chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_id=f"{document_id}_chunk_{chunk_id}",
                    content=chunk_content,
                    chunk_type=chunk_type,
                    word_count=len(chunk_content.split())
                )
                chunks.append(chunk)
                chunk_id += 1
                
        return chunks
    
    def classify_chunk_type(self, text: str) -> str:
        """Classify the type of content in a chunk."""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['instruction', 'step', 'procedure', 'how to']):
            return 'instructions'
        elif any(word in text_lower for word in ['rule', 'policy', 'requirement', 'must', 'should']):
            return 'rules'
        elif any(word in text_lower for word in ['context', 'background', 'overview', 'about']):
            return 'context'
        else:
            return 'general'
    
    def analyze_chunk_with_llm(self, chunk: DocumentChunk) -> Optional[AnalysisResult]:
        """Analyze a single chunk using LLM."""
        prompt = self.create_analysis_prompt(chunk)
        
        try:
            # Prioritize Anthropic Claude
            if self.anthropic_client:
                response = self.anthropic_client.messages.create(
                    model=self.config.get('anthropic_model', 'claude-3-sonnet-20240229'),
                    max_tokens=2000,
                    temperature=0.3,
                    messages=[{"role": "user", "content": prompt}]
                )
                analysis_text = response.content[0].text
            elif self.openai_client:
                from openai import OpenAI
                client = OpenAI(api_key=self.config['openai_api_key'])
                response = client.chat.completions.create(
                    model=self.config.get('openai_model', 'gpt-4'),
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=2000
                )
                analysis_text = response.choices[0].message.content
            else:
                self.logger.error("No LLM client configured")
                return None
            
            # Store raw response for this chunk
            self.store_raw_response(chunk, analysis_text)
                
            return self.parse_analysis_response(chunk, analysis_text)
            
        except Exception as e:
            self.logger.error(f"LLM analysis failed for chunk {chunk.chunk_id}: {e}")
            return None
    
    def store_raw_response(self, chunk: DocumentChunk, response: str) -> None:
        """Store raw LLM response for a chunk."""
        document_id = chunk.document_id
        
        if document_id not in self.raw_responses:
            self.raw_responses[document_id] = []
        
        self.raw_responses[document_id].append({
            'chunk_id': chunk.chunk_id,
            'chunk_type': chunk.chunk_type,
            'content': chunk.content,
            'response': response,
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    def create_analysis_prompt(self, chunk: DocumentChunk) -> str:
        """Create a prompt for analyzing a document chunk."""
        return f"""
You are analyzing merchant instruction documents to extract generalizable rules and patterns.

Document Chunk Type: {chunk.chunk_type}
Content:
{chunk.content}

Please analyze this content and provide:

1. **Core Themes** (3-5 main themes/topics):
   - List the main themes or topics covered

2. **Extracted Rules** (specific rules or requirements):
   - List any specific rules, requirements, or policies mentioned
   - Include both explicit and implicit rules

3. **Exceptions or Special Cases**:
   - Note any exceptions, special cases, or conditional requirements

4. **Merchant-Specific Elements**:
   - Identify elements that might need customization per merchant
   - Note any merchant-specific terminology or processes

5. **Confidence Score** (0.0 to 1.0):
   - Rate your confidence in the analysis

Format your response as JSON with these exact keys: themes, rules, exceptions, merchant_specific, confidence_score
"""
    
    def parse_analysis_response(self, chunk: DocumentChunk, response: str) -> AnalysisResult:
        """Parse LLM response into structured data."""
        try:
            # Extract JSON from response
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                data = json.loads(json_match.group())
            else:
                # Fallback parsing if JSON not found
                data = self.fallback_parse(response)
                
            return AnalysisResult(
                chunk_id=chunk.chunk_id,
                document_id=chunk.document_id,
                themes=data.get('themes', []),
                rules=data.get('rules', []),
                exceptions=data.get('exceptions', []),
                merchant_specific=data.get('merchant_specific', []),
                confidence_score=float(data.get('confidence_score', 0.5)),
                analysis_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
            )
        except Exception as e:
            self.logger.error(f"Failed to parse analysis response: {e}")
            return AnalysisResult(
                chunk_id=chunk.chunk_id,
                document_id=chunk.document_id,
                themes=[],
                rules=[],
                exceptions=[],
                merchant_specific=[],
                confidence_score=0.0,
                analysis_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
            )
    
    def fallback_parse(self, response: str) -> Dict[str, Any]:
        """Fallback parsing when JSON extraction fails."""
        data = {
            'themes': [],
            'rules': [],
            'exceptions': [],
            'merchant_specific': [],
            'confidence_score': 0.5
        }
        
        # Simple text parsing as fallback
        lines = response.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if 'theme' in line.lower():
                current_section = 'themes'
            elif 'rule' in line.lower():
                current_section = 'rules'
            elif 'exception' in line.lower():
                current_section = 'exceptions'
            elif 'merchant' in line.lower():
                current_section = 'merchant_specific'
            elif line.startswith('-') or line.startswith('•'):
                if current_section and current_section in data:
                    data[current_section].append(line[1:].strip())
                    
        return data
    
    def synthesize_rules(self) -> List[SynthesizedRule]:
        """Synthesize rules from all analysis results."""
        # Collect all rules
        all_rules = []
        for result in self.analysis_results:
            for rule in result.rules:
                # Handle both string and dict rules
                if isinstance(rule, dict):
                    rule_text = rule.get('rule', str(rule))
                else:
                    rule_text = str(rule)
                
                all_rules.append({
                    'rule': rule_text,
                    'source_document': result.document_id,
                    'confidence': result.confidence_score
                })
        
        # Group similar rules
        rule_groups = self.group_similar_rules(all_rules)
        
        # Create synthesized rules
        synthesized = []
        for i, group in enumerate(rule_groups):
            if len(group) < 2:  # Skip rules that appear only once
                continue
                
            # Find the most common formulation
            rule_texts = [str(item['rule']) for item in group]
            self.logger.debug(f"Rule texts for group {i}: {rule_texts[:2]}")  # Debug first 2
            most_common = Counter(rule_texts).most_common(1)[0][0]
            
            # Extract merchant customizations
            merchant_customizations = {}
            for item in group:
                doc_id = item['source_document']
                if doc_id not in merchant_customizations:
                    merchant_customizations[doc_id] = item['rule']
            
            synthesized_rule = SynthesizedRule(
                rule_id=f"rule_{i}",
                general_rule=most_common,
                merchant_customizations=merchant_customizations,
                frequency=len(group),
                confidence=np.mean([item['confidence'] for item in group]),
                source_documents=list(set([item['source_document'] for item in group]))
            )
            synthesized.append(synthesized_rule)
            
        return synthesized
    
    def group_similar_rules(self, rules: List[Dict]) -> List[List[Dict]]:
        """Group similar rules together using simple text similarity."""
        groups = []
        used_indices = set()
        
        for i, rule1 in enumerate(rules):
            if i in used_indices:
                continue
                
            group = [rule1]
            used_indices.add(i)
            
            for j, rule2 in enumerate(rules[i+1:], i+1):
                if j in used_indices:
                    continue
                    
                if self.rules_similar(rule1['rule'], rule2['rule']):
                    group.append(rule2)
                    used_indices.add(j)
                    
            groups.append(group)
            
        return groups
    
    def rules_similar(self, rule1: str, rule2: str) -> bool:
        # Ensure we have strings
        if not isinstance(rule1, str) or not isinstance(rule2, str):
            return False
        """Check if two rules are similar using simple text similarity."""
        # Simple similarity based on common words
        words1 = set(rule1.lower().split())
        words2 = set(rule2.lower().split())
        
        # Remove common stop words
        try:
            stop_words = set(stopwords.words('english'))
        except:
            stop_words = set(['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'])
        words1 = words1 - stop_words
        words2 = words2 - stop_words
        
        if not words1 or not words2:
            return False
            
        # Calculate Jaccard similarity
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        
        similarity = intersection / union if union > 0 else 0
        return similarity > self.config.get("similarity_threshold", 0.6)  # Threshold for similarity
    
    def process_pdfs(self, pdf_directory: str) -> None:
        """Process all PDFs in a directory."""
        pdf_path = Path(pdf_directory)
        if not pdf_path.exists():
            raise ValueError(f"Directory {pdf_directory} does not exist")
            
        pdf_files = list(pdf_path.glob("*.pdf"))
        docx_files = list(pdf_path.glob("*.docx"))
        all_files = pdf_files + docx_files
        self.logger.info(f"Found {len(pdf_files)} PDF files and {len(docx_files)} DOCX files to process")
        
        # Process PDFs in batches
        batch_size = self.config.get('batch_size', 10)
        for i in range(0, len(all_files), batch_size):
            batch = all_files[i:i + batch_size]
            self.logger.info(f"Processing batch {i//batch_size + 1}/{(len(pdf_files) + batch_size - 1)//batch_size}")
            
            for file_path in batch:
                try:
                    self.process_single_pdf(file_path)
                except Exception as e:
                    self.logger.error(f"Failed to process {file_path}: {e}")
    
    def process_single_pdf(self, pdf_path: Path) -> None:
        """Process a single PDF file."""
        document_id = pdf_path.stem
        
        # Extract text
        self.logger.info(f"Extracting text from {pdf_path}")
        text = self.extract_text_from_file(str(pdf_path))
        
        if not text.strip():
            self.logger.warning(f"No text extracted from {pdf_path}")
            return
            
        # Preprocess text
        text = self.preprocess_text(text)
        
        # Chunk document
        chunks = self.chunk_document(document_id, text)
        self.logger.info(f"Created {len(chunks)} chunks from {pdf_path}")
        
        # Store chunks
        self.chunks.extend(chunks)
        
        # Analyze chunks with LLM
        self.analyze_chunks_batch(chunks)
    
    def analyze_chunks_batch(self, chunks: List[DocumentChunk]) -> None:
        """Analyze a batch of chunks using LLM."""
        max_workers = self.config.get('max_workers', 5)
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_chunk = {
                executor.submit(self.analyze_chunk_with_llm, chunk): chunk 
                for chunk in chunks
            }
            
            for future in as_completed(future_to_chunk):
                chunk = future_to_chunk[future]
                try:
                    result = future.result()
                    if result:
                        self.analysis_results.append(result)
                        self.logger.info(f"Analyzed chunk {chunk.chunk_id}")
                except Exception as e:
                    self.logger.error(f"Analysis failed for chunk {chunk.chunk_id}: {e}")
    
    def generate_report(self, output_dir: str) -> None:
        """Generate comprehensive analysis report."""
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # Synthesize rules
        self.logger.info("Synthesizing rules...")
        self.synthesized_rules = self.synthesize_rules()
        
        # Generate JSON report
        report_data = {
            'summary': {
                'total_documents': len(set([chunk.document_id for chunk in self.chunks])),
                'total_chunks': len(self.chunks),
                'total_analysis_results': len(self.analysis_results),
                'synthesized_rules': len(self.synthesized_rules)
            },
            'synthesized_rules': [asdict(rule) for rule in self.synthesized_rules],
            'analysis_results': [asdict(result) for result in self.analysis_results]
        }
        
        with open(output_path / 'analysis_report.json', 'w') as f:
            json.dump(report_data, f, indent=2)
        
        # Generate human-readable report
        self.generate_human_report(output_path)
        
        # Generate CSV for further analysis
        self.generate_csv_reports(output_path)
        
        # Generate raw response files
        self.generate_raw_response_files(output_path)
        
        self.logger.info(f"Report generated in {output_path}")
    
    def generate_human_report(self, output_path: Path) -> None:
        """Generate a human-readable markdown report."""
        report_content = f"""# PDF Rule Analysis Report

## Summary
- **Total Documents Processed**: {len(set([chunk.document_id for chunk in self.chunks]))}
- **Total Chunks Analyzed**: {len(self.chunks)}
- **Synthesized Rules**: {len(self.synthesized_rules)}

## Generalizable Rules

"""
        
        for rule in self.synthesized_rules:
            report_content += f"""### Rule {rule.rule_id}
**General Rule**: {rule.general_rule}

**Frequency**: Appears in {rule.frequency} documents
**Confidence**: {rule.confidence:.2f}

**Merchant-Specific Variations**:
"""
            for merchant, variation in rule.merchant_customizations.items():
                report_content += f"- **{merchant}**: {variation}\n"
            
            report_content += f"\n**Source Documents**: {', '.join(rule.source_documents)}\n\n"
        
        with open(output_path / 'analysis_report.md', 'w') as f:
            f.write(report_content)
    
    def generate_csv_reports(self, output_path: Path) -> None:
        """Generate CSV files for further analysis."""
        # Rules CSV
        rules_data = []
        for rule in self.synthesized_rules:
            rules_data.append({
                'rule_id': rule.rule_id,
                'general_rule': rule.general_rule,
                'frequency': rule.frequency,
                'confidence': rule.confidence,
                'source_documents': '; '.join(rule.source_documents)
            })
        
        pd.DataFrame(rules_data).to_csv(output_path / 'synthesized_rules.csv', index=False)
        
        # Analysis results CSV
        analysis_data = []
        for result in self.analysis_results:
            analysis_data.append({
                'chunk_id': result.chunk_id,
                'document_id': result.document_id,
                'themes': '; '.join([str(t) for t in result.themes]),
                'rules': '; '.join([str(r) for r in result.rules]),
                'exceptions': '; '.join([str(e) for e in result.exceptions]),
                'merchant_specific': '; '.join([str(m) for m in result.merchant_specific]),
                'confidence_score': result.confidence_score
            })
        
        pd.DataFrame(analysis_data).to_csv(output_path / 'analysis_results.csv', index=False)
    
    def generate_raw_response_files(self, output_path: Path) -> None:
        """Generate individual markdown files for each document's raw LLM responses."""
        raw_responses_dir = output_path / 'raw_responses'
        raw_responses_dir.mkdir(exist_ok=True)
        
        for document_id, responses in self.raw_responses.items():
            # Clean document name for filename
            safe_filename = re.sub(r'[^\w\s-]', '', document_id).strip()
            safe_filename = re.sub(r'[-\s]+', '_', safe_filename)
            
            markdown_content = f"""# Raw LLM Analysis: {document_id}

## Document Overview
- **Total Chunks Analyzed**: {len(responses)}
- **Analysis Timestamp**: {responses[0]['timestamp'] if responses else 'N/A'}

---

"""
            
            for i, response_data in enumerate(responses, 1):
                markdown_content += f"""## Chunk {i}: {response_data['chunk_id']}

**Chunk Type**: {response_data['chunk_type']}  
**Analysis Time**: {response_data['timestamp']}

### Original Content
```
{response_data['content'][:500]}{'...' if len(response_data['content']) > 500 else ''}
```

### LLM Analysis Response
{response_data['response']}

---

"""
            
            # Write markdown file
            filename = f"{safe_filename}.md"
            with open(raw_responses_dir / filename, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            self.logger.info(f"Raw responses saved for {document_id} -> {filename}")

def main():
    """Main function to run the PDF rule extractor."""
    parser = argparse.ArgumentParser(description='Extract generalizable rules from PDF documents')
    parser.add_argument('--pdf-dir', required=True, help='Directory containing PDF files')
    parser.add_argument('--output-dir', required=True, help='Output directory for results')
    parser.add_argument('--config', help='Configuration file path')
    parser.add_argument('--openai-key', help='OpenAI API key')
    parser.add_argument('--anthropic-key', help='Anthropic API key')
    
    args = parser.parse_args()
    
    # Load configuration
    config = {
        'batch_size': 10,
        'max_workers': 5,
        'max_sentences_per_chunk': 10,
        'openai_model': 'gpt-4',
        'anthropic_model': 'claude-3-sonnet-20240229'
    }
    
    if args.config:
        with open(args.config, 'r') as f:
            config.update(json.load(f))
    
    if args.openai_key:
        config['openai_api_key'] = args.openai_key
    
    if args.anthropic_key:
        config['anthropic_api_key'] = args.anthropic_key
    
    # Create extractor and process PDFs
    extractor = PDFRuleExtractor(config)
    
    try:
        extractor.process_pdfs(args.pdf_dir)
        extractor.generate_report(args.output_dir)
        print(f"Analysis complete! Results saved to {args.output_dir}")
    except Exception as e:
        print(f"Error: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())
